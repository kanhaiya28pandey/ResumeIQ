import io
import pdfplumber
from docx import Document
from fastapi import HTTPException

def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)

def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    text_parts = [para.text for para in doc.paragraphs if para.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    return "\n".join(text_parts)

def extract_resume_text(filename: str, file_bytes: bytes) -> str:
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    else:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
    if not text or len(text.strip()) < 20:
        raise HTTPException(
            status_code=400,
            detail="Could not extract text from this file. Make sure it's not a scanned image.",
        )

    return text.strip()